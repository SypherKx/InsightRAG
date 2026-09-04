"""
Document ingestion module for RAG pipeline.

Handles loading and extracting text from various file formats:
- PDF (.pdf)
- Plain text (.txt)
- Markdown (.md)
- Future: DOCX, HTML, etc.
"""

import os
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
import logging

from .models import Document

logger = logging.getLogger(__name__)


@dataclass
class IngestionConfig:
    """Configuration for document ingestion."""
    allowed_extensions: List[str] = None
    max_file_size_mb: int = 50
    recursive: bool = False  # Recursively scan directories
    encoding: str = "utf-8"

    def __post_init__(self):
        if self.allowed_extensions is None:
            self.allowed_extensions = [
                ".pdf", ".txt", ".md", ".csv", ".json", ".log",
                ".rst", ".html", ".xml", ".docx",
                ".png", ".jpg", ".jpeg", ".webp", ".bmp"
            ]


class DocumentIngester:
    """
    Handles document loading and text extraction.

    Supports multiple file formats with proper error handling.
    Extracts text content and basic metadata.
    """

    def __init__(self, config: Optional[IngestionConfig] = None):
        """
        Initialize document ingester.

        Args:
            config: Ingestion configuration
        """
        self.config = config or IngestionConfig()
        self._loaded_extractors = self._init_extractors()

    def _init_extractors(self) -> Dict[str, callable]:
        """Initialize file format extractors."""
        return {
            ".pdf": self._extract_pdf,
            ".txt": self._extract_txt,
            ".md": self._extract_markdown,
            ".csv": self._extract_csv,
            ".json": self._extract_json,
            ".log": self._extract_txt,
            ".rst": self._extract_txt,
            ".html": self._extract_txt,
            ".xml": self._extract_txt,
            ".docx": self._extract_docx,
            ".png": self._extract_image,
            ".jpg": self._extract_image,
            ".jpeg": self._extract_image,
            ".webp": self._extract_image,
            ".bmp": self._extract_image,
        }

    def ingest_file(self, file_path: Union[str, Path], org_id: str,
                    document_type: Optional[str] = None,
                    start_page: Optional[int] = None,
                    end_page: Optional[int] = None) -> Optional[Document]:
        """
        Ingest a single file with optional page range slicing.

        Args:
            file_path: Path to file
            org_id: Organization ID for multi-tenancy
            document_type: Type classification (auto-detected if None)
            start_page: Optional 1-indexed starting page
            end_page: Optional 1-indexed ending page

        Returns:
            Document object or None if failed
        """
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        # Check file extension
        ext = file_path.suffix.lower()
        if ext not in self.config.allowed_extensions:
            logger.warning(f"Unsupported file type: {ext} for {file_path}")
            return None

        # Check file size
        size_mb = file_path.stat().st_size / (1024 * 1024)
        if size_mb > self.config.max_file_size_mb:
            logger.error(f"File too large: {size_mb:.1f}MB > {self.config.max_file_size_mb}MB")
            return None

        # Auto-detect document type from path or filename if not provided
        if document_type is None:
            document_type = self._detect_document_type(file_path)

        try:
            # Extract text based on file type
            if ext == ".pdf":
                content, metadata = self._extract_pdf(file_path, start_page=start_page, end_page=end_page)
            elif ext in self._loaded_extractors:
                content, metadata = self._loaded_extractors[ext](file_path)
            else:
                logger.error(f"No extractor for extension: {ext}")
                return None

            if not content or not content.strip():
                logger.warning(f"No content extracted from {file_path}")
                return None

            # Create Document object
            doc = Document(
                id=str(uuid.uuid4()),
                org_id=org_id,
                title=file_path.stem,
                source_path=str(file_path),
                document_type=document_type,
                content=content,
                metadata={
                    **metadata,
                    "file_size_bytes": file_path.stat().st_size,
                    "file_extension": ext,
                    "file_name": file_path.name,
                    "start_page": start_page,
                    "end_page": end_page,
                }
            )

            range_str = f" (Pages {start_page or 1}-{end_page or metadata.get('page_count', 1)})" if (start_page or end_page) else ""
            logger.info(f"Ingested {file_path}{range_str} ({len(content)} chars, {metadata.get('page_count', 0)} total pages)")
            return doc

        except Exception as e:
            logger.exception(f"Failed to ingest {file_path}: {e}")
            return None

    def ingest_directory(self, directory_path: Union[str, Path], org_id: str,
                         document_type: Optional[str] = None) -> List[Document]:
        """
        Ingest all supported files in a directory.

        Args:
            directory_path: Directory to scan
            org_id: Organization ID
            document_type: Override document type (auto-detected per file if None)

        Returns:
            List of successfully ingested documents
        """
        directory_path = Path(directory_path)

        if not directory_path.is_dir():
            logger.error(f"Not a directory: {directory_path}")
            return []

        documents = []
        pattern = "**/*" if self.config.recursive else "*"

        for file_path in directory_path.glob(pattern):
            if file_path.is_file():
                ext = file_path.suffix.lower()
                if ext in self.config.allowed_extensions:
                    doc = self.ingest_file(file_path, org_id, document_type)
                    if doc:
                        documents.append(doc)

        logger.info(f"Ingested {len(documents)} documents from {directory_path}")
        return documents

    def _detect_document_type(self, file_path: Path) -> str:
        """
        Detect document type from filename or path.

        Args:
            file_path: File path

        Returns:
            Document type string
        """
        name = file_path.stem.lower()

        # Check for known keywords in filename
        type_mappings = {
            "metric": "metric_def",
            "definition": "metric_def",
            "glossary": "metric_def",
            "process": "process",
            "procedure": "process",
            "runbook": "process",
            "incident": "incident",
            "postmortem": "incident",
            "outage": "incident",
            "issue": "incident",
        }

        for keyword, doc_type in type_mappings.items():
            if keyword in name:
                return doc_type

        # Check directory path
        parent = file_path.parent.name.lower()
        if "process" in parent or "procedure" in parent:
            return "process"
        if "incident" in parent or "postmortem" in parent:
            return "incident"
        if "metric" in parent or "definition" in parent:
            return "metric_def"

        # Default
        return "other"

    def _extract_pdf(self, file_path: Path, start_page: Optional[int] = None, end_page: Optional[int] = None) -> tuple[str, dict]:
        """
        Extract text from PDF file using PyMuPDF (fitz), pypdf, or PyPDF2 with optional page range.

        Args:
            file_path: PDF file path
            start_page: 1-indexed start page
            end_page: 1-indexed end page

        Returns:
            Tuple: (text_content, metadata)
        """
        pages = []
        metadata = {
            "page_count": 0,
            "author": None,
            "title": None,
            "creation_date": None,
        }

        # 1. Try PyMuPDF (fitz) - fastest & highest fidelity
        try:
            import fitz
            doc = fitz.open(str(file_path))
            total_pages = len(doc)
            metadata["page_count"] = total_pages
            if doc.metadata:
                metadata["author"] = doc.metadata.get("author")
                metadata["title"] = doc.metadata.get("title")
                metadata["creation_date"] = doc.metadata.get("creationDate")

            s_idx = max(0, (start_page - 1)) if start_page else 0
            e_idx = min(total_pages, end_page) if end_page else total_pages

            for p_num in range(s_idx, e_idx):
                page = doc[p_num]
                text = page.get_text()
                if text and text.strip():
                    pages.append(f"[Page {p_num + 1}]\n" + text.strip())
            doc.close()
            if pages:
                metadata["page_range"] = f"Pages {s_idx + 1}-{e_idx}"
                return "\n\n".join(pages), metadata
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"fitz PDF extraction failed for {file_path}: {e}")

        # 2. Try PyPDF2
        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                total_pages = len(reader.pages)
                metadata["page_count"] = total_pages
                s_idx = max(0, (start_page - 1)) if start_page else 0
                e_idx = min(total_pages, end_page) if end_page else total_pages

                for p_num in range(s_idx, e_idx):
                    text = reader.pages[p_num].extract_text()
                    if text and text.strip():
                        pages.append(f"[Page {p_num + 1}]\n" + text.strip())
            if pages:
                metadata["page_range"] = f"Pages {s_idx + 1}-{e_idx}"
                return "\n\n".join(pages), metadata
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed for {file_path}: {e}")

        # 3. Try pypdf
        try:
            import pypdf
            reader = pypdf.PdfReader(str(file_path))
            total_pages = len(reader.pages)
            metadata["page_count"] = total_pages
            s_idx = max(0, (start_page - 1)) if start_page else 0
            e_idx = min(total_pages, end_page) if end_page else total_pages

            for p_num in range(s_idx, e_idx):
                text = reader.pages[p_num].extract_text()
                if text and text.strip():
                    pages.append(f"[Page {p_num + 1}]\n" + text.strip())
            if pages:
                metadata["page_range"] = f"Pages {s_idx + 1}-{e_idx}"
                return "\n\n".join(pages), metadata
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"pypdf extraction failed for {file_path}: {e}")

        content = "\n\n".join(pages)
        if not content:
            raise ValueError(f"Could not extract readable text from PDF: {file_path}")
        return content, metadata

    def _extract_txt(self, file_path: Path) -> tuple[str, dict]:
        """
        Extract text from plain text file.

        Args:
            file_path: Text file path

        Returns:
            Tuple: (text_content, metadata)
        """
        content = None
        used_enc = self.config.encoding
        for encoding in [self.config.encoding, 'utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
            try:
                with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                    content = f.read()
                used_enc = encoding
                break
            except Exception:
                continue

        if content is None:
            raise ValueError(f"Unable to decode text file: {file_path}")

        metadata = {
            "file_size": file_path.stat().st_size,
            "encoding": used_enc,
        }
        return content, metadata

    def _extract_markdown(self, file_path: Path) -> tuple[str, dict]:
        """Extract text from Markdown file."""
        content, metadata = self._extract_txt(file_path)
        metadata["file_type"] = "markdown"
        return content, metadata

    def _extract_csv(self, file_path: Path) -> tuple[str, dict]:
        """Extract text from CSV file with column structure."""
        try:
            import csv
            rows_summary = []
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                reader = csv.reader(f)
                headers = next(reader, None)
                if headers:
                    rows_summary.append(f"Headers: {', '.join(headers)}")
                    for idx, row in enumerate(reader):
                        if idx < 500:  # Cap at first 500 rows for concise RAG context
                            item_strs = [f"{h}: {v}" for h, v in zip(headers, row) if v]
                            rows_summary.append(f"Row {idx+1}: {'; '.join(item_strs)}")
                        else:
                            break
            content = "\n".join(rows_summary)
            return content, {"file_type": "csv", "file_size": file_path.stat().st_size}
        except Exception:
            return self._extract_txt(file_path)

    def _extract_json(self, file_path: Path) -> tuple[str, dict]:
        """Extract text from JSON file."""
        try:
            import json
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                data = json.load(f)
            formatted = json.dumps(data, indent=2)
            return formatted, {"file_type": "json", "file_size": file_path.stat().st_size}
        except Exception:
            return self._extract_txt(file_path)

    def _extract_docx(self, file_path: Path) -> tuple[str, dict]:
        """Extract text from DOCX file with decompression bomb defense."""
        try:
            import docx
            doc = docx.Document(str(file_path))
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paras)
            return content, {"file_type": "docx", "paragraphs_count": len(paras)}
        except Exception:
            # Fallback using zipfile xml extraction with safe size checks
            import zipfile
            import xml.etree.ElementTree as ET
            with zipfile.ZipFile(file_path) as z:
                # Zip bomb defense: check uncompressed size before reading
                info = z.getinfo('word/document.xml')
                if info.file_size > 50 * 1024 * 1024:  # Max 50MB XML uncompressed
                    raise ValueError("DOCX document XML exceeds maximum safe size.")
                xml_content = z.read('word/document.xml')
            tree = ET.fromstring(xml_content)
            texts = [node.text for node in tree.iter() if node.text]
            return " ".join(texts), {"file_type": "docx"}

    def _extract_image(self, file_path: Path) -> tuple[str, dict]:
        """Extract text/metadata from image file with decompression bomb defense."""
        metadata = {
            "file_type": "image",
            "file_size": file_path.stat().st_size,
            "page_count": 1,
            "page": 1,
        }
        text = f"Visual Document & Diagram: {file_path.name}"
        try:
            from PIL import Image
            # Protect against pixel flood / decompression bomb DoS
            Image.MAX_IMAGE_PIXELS = 25_000_000
            with Image.open(file_path) as img:
                metadata["width"] = img.width
                metadata["height"] = img.height
                metadata["format"] = img.format or file_path.suffix.lstrip(".")
                text = f"Visual Diagram File: {file_path.name} (Image Resolution: {img.width}x{img.height} px)."
        except Exception as e:
            logger.warning(f"Could not open image for metadata inspection: {e}")

        try:
            import pytesseract
            from PIL import Image
            Image.MAX_IMAGE_PIXELS = 25_000_000
            ocr_text = pytesseract.image_to_string(Image.open(file_path)).strip()
            if ocr_text:
                text += f"\n\nOCR Extracted Text Content:\n{ocr_text}"
        except Exception:
            pass

        return text, metadata


def create_default_ingester() -> DocumentIngester:
    """Create ingester with default configuration."""
    config = IngestionConfig(
        allowed_extensions=[
            ".pdf", ".txt", ".md", ".csv", ".json", ".log",
            ".rst", ".html", ".xml", ".docx",
            ".png", ".jpg", ".jpeg", ".webp", ".bmp"
        ],
        max_file_size_mb=50,
        recursive=False
    )
    return DocumentIngester(config)
